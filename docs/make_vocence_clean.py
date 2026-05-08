"""Build vocence-clean.docx — task brief without protocol/payment terminology."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_kv(doc, k, v):
    p = doc.add_paragraph()
    r1 = p.add_run(k + ": ")
    r1.bold = True
    p.add_run(v)
    return p


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            t.rows[r_i].cells[c_i].text = str(val)
    return t


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def main():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    add_h1(doc, "Vocence: Prompt-Conditioned Text-to-Speech — Task Brief")

    add_kv(doc, "Base model",
           "https://huggingface.co/Qwen/Qwen3-TTS-12Hz-1.7B-VoiceDesign")
    add_kv(doc, "Training dataset",
           "LibriVox audiobook clips (public-domain, CC0)")
    add_kv(doc, "Output sample rate", "24,000 Hz mono WAV")

    add_h2(doc, "Overview")
    add_p(doc,
          "Build a text-to-speech engine that takes a text passage and a "
          "style instruction (gender, age group, pitch, speed, emotion, "
          "tone, accent) and produces a single audio clip that (a) pronounces "
          "the text correctly, (b) matches every style trait, and (c) sounds "
          "as natural as a real human recording of the same passage in the "
          "same style. The hard problem is the third axis — perceptual "
          "naturalness against a real human voice.")

    add_h2(doc, "Task Definition")
    add_p(doc,
          "Each task consists of one text passage (up to 2000 characters) "
          "and one structured style instruction (up to 600 characters). The "
          "engine must return a 24 kHz mono waveform within a 300-second wall "
          "budget. Each output is independently scored on 9 elements; a "
          "weighted aggregate must clear a 0.9 pass threshold to count "
          "toward the run.")

    add_h2(doc, "Engine Contract")
    add_table(doc, ["Constraint", "Specification"], [
        ["Class API",        "Engine(path_hf_repo)  with .warmup() and "
                             ".generate_wav(instruction, text) -> (np.ndarray, int)"],
        ["Imports",          "Standard library + site-packages only"],
        ["External network", "Prohibited — engine must run fully local"],
        ["Sample rate",      "24,000 Hz"],
        ["Max text length",  "2,000 characters"],
        ["Max instruction",  "600 characters"],
        ["Per-call deadline","300 seconds wall time"],
        ["Container",        "Single-GPU image, ≥48 GB VRAM"],
    ])

    add_h2(doc, "Evaluation Pipeline (Per Sample)")
    add_p(doc,
          "Each evaluation triple is (text, instruction, source_clip), where "
          "the source clip is a real LibriVox recording of the same text by a "
          "voice that matches the instruction. The judge runs two GPT-4o-audio "
          "calls:")
    add_bullet(doc,
               "Pointwise extraction — given only the engine's output, the "
               "judge produces a transcription and predicts each of the 7 "
               "categorical traits.")
    add_bullet(doc,
               "Pairwise naturalness — given the engine's output and the "
               "source clip in randomized order, the judge picks which one "
               "sounds more natural (clarity, prosody, intonation, absence "
               "of synthetic artifacts).")
    add_p(doc, "These signals are combined into the per-element score table:")

    add_table(doc, ["Element", "Source", "Weight"], [
        ["script (WER vs target)",          "transcription from pointwise",     "30%"],
        ["naturalness (pairwise outcome)",  "FIRST/SECOND from pairwise",        "15%"],
        ["gender",                          "predicted vs spec",                 "10%"],
        ["speed",                           "predicted vs spec (ordinal)",       "10%"],
        ["emotion",                         "predicted vs spec",                 "10%"],
        ["age_group",                       "predicted vs spec (ordinal)",       "10%"],
        ["pitch",                           "predicted vs spec (ordinal)",       "5%"],
        ["accent",                          "predicted vs spec",                 "5%"],
        ["tone",                            "predicted vs spec",                 "5%"],
    ])
    add_p(doc, "Aggregate must clear 0.9 to mark a pass on that sample. "
               "Ordinal traits use a 1.0 / 0.5 / 0.0 distance scoring; "
               "categorical traits are binary match. Naturalness is binary "
               "(1.0 if the judge prefers the engine output over the source, "
               "else 0.0).")

    add_h2(doc, "Selection Logic")
    add_p(doc,
          "Each round, all submissions are scored on the same task batch. "
          "The submission with the highest pass rate over the batch holds the "
          "top position. Ties and near-ties resolve by a paired comparison "
          "test on per-sample scores. New submissions only displace the "
          "current top if they consistently win that test across multiple "
          "rounds — a single lucky batch is not enough.")

    add_h2(doc, "Current Competitive Landscape")
    add_table(doc, ["Metric", "Value"], [
        ["Top pass rate to beat",        "~63%"],
        ["Plausible legitimate ceiling", "70-75%"],
        ["Per-sample budget",            "300 s wall, ~30 s typical"],
        ["Active engines",               "Several, mostly Qwen3-TTS-VoiceDesign or Parler-TTS-mini variants with SFT on LibriVox"],
    ])

    add_h2(doc, "Where the Difficulty Sits")
    add_bullet(doc,
               "Script (transcription / WER) routinely scores >0.98. Trait "
               "matching (gender, speed, emotion, age, pitch, accent, tone) "
               "scores 0.85-1.0. These are not the bottleneck.")
    add_bullet(doc,
               "Naturalness scores 0.50-0.65 for well-tuned engines. The "
               "judge is comparing TTS output to a real human audiobook "
               "reading; the perceptual gap on prosody, breath, and "
               "micro-timing is what costs aggregate points.")
    add_bullet(doc,
               "A 1pp lift on naturalness yields ~0.0015 of aggregate, "
               "which is enough to separate two otherwise-similar engines.")

    add_h2(doc, "Working Recipe (Established Techniques)")

    add_p(doc, "1. Supervised fine-tune the base on LibriVox text/audio pairs.")
    add_bullet(doc, "3 epochs, constant lr=2e-6, bf16, gradient_accum=4 across 4 GPUs.")
    add_bullet(doc, "Speaker-level (book-level) train/heldout split prevents leakage.")
    add_bullet(doc, "Filter clips by perceptual MOS (UTMOSv2) to remove the bottom quartile of "
                    "amateur or noisy readers — keeps prosodic diversity in the kept set.")

    add_p(doc, "2. Best-of-N inference with composite picker.")
    add_bullet(doc, "Generate K=8 candidates per task using a temperature ladder "
                    "(T = 0.6, 0.7, 0.8, 0.9, 1.0). Sampling diversity matters more than any single setting.")
    add_bullet(doc, "Pick the best candidate via 0.3 × UTMOSv2 + 0.7 × (1 − faster-whisper-WER).")
    add_bullet(doc, "Use Qwen3-TTS's built-in batched generation for ~6× wall-time speedup over sequential sampling.")

    add_p(doc, "3. Train a small picker that mimics the GPT-4o-audio judge.")
    add_bullet(doc, "Collect ~600-3000 (candidate, source, judgment) triples by running the judge "
                    "on the engine's own outputs (cost: ~$0.04 per pair).")
    add_bullet(doc, "Distill into a WavLM-base + 2-layer MLP that predicts pairwise preference.")
    add_bullet(doc, "Replace the composite picker with this distilled model; lift is typically 5-8pp on naturalness.")

    add_p(doc, "4. Direct Preference Optimization with model-generated pairs.")
    add_bullet(doc, "Critical: use the engine's own samples for both 'chosen' and 'rejected' sides. "
                    "Using the LibriVox source as 'chosen' degrades performance (Koel-TTS, ICML 2025).")
    add_bullet(doc, "5,000-10,000 pairs is the sweet spot; >50,000 has diminishing returns.")
    add_bullet(doc, "RPO (score-aware preference optimization) is preferred over vanilla DPO — "
                    "scales gradient by the score gap, less sensitive to β.")
    add_bullet(doc, "Hyperparameters: β=0.01, lr=2e-7, batch=64 pairs, ~4000 steps. "
                    "Iterate two rounds — generate fresh pairs from the latest checkpoint each time.")

    add_p(doc, "5. Classifier-Free Guidance at inference.")
    add_bullet(doc, "Drop the style instruction at runtime to obtain unconditional logits, "
                    "then interpolate: logit_final = γ · logit_cond + (1 − γ) · logit_uncond, "
                    "with γ ≈ 2.5.")
    add_bullet(doc, "If the architecture doesn't natively expose CFG, run a 500-step continued "
                    "SFT pass with 10% conditioning dropout — a few hours of GPU time. "
                    "CFG alone has been documented to add 5-10pp on naturalness.")

    add_h2(doc, "Things That Don't Work")
    add_bullet(doc, "Loudness normalization (LUFS) and light bandwidth shaping (lowpass at 12 kHz). "
                    "The judge picks up clarity and prosody, not spectral envelope. Neutral effect.")
    add_bullet(doc, "Light reverb / room tone injection. Too subtle to register; too obvious to be safe.")
    add_bullet(doc, "Filtering training data by MOS alone (no other diversity criterion) — "
                    "the model converges to a flat 'safe' voice.")
    add_bullet(doc, "Treating the source recording as the SFT or DPO target. The base model "
                    "already approximates LibriVox; the marginal lift comes from preference learning, "
                    "not more imitation.")

    add_h2(doc, "Reference Material")
    add_bullet(doc, "Koel-TTS (NVIDIA, ICML 2025) — DPO/RPO recipe + CFG hyperparameters: arXiv:2502.05236")
    add_bullet(doc, "SpeechAlign (2024) — iterative DPO for speech token models: arXiv:2404.05600")
    add_bullet(doc, "GSRM (Generative Speech Quality Model) — distill audio judge into small classifier: arXiv:2602.13891")
    add_bullet(doc, "UTMOSv2 — perceptual MOS predictor: github.com/sarulab-speech/UTMOSv2")
    add_bullet(doc, "NISQA-TTS — naturalness MOS head: github.com/gabrielmittag/NISQA")
    add_bullet(doc, "Qwen3-TTS upstream: github.com/QwenLM/Qwen3-TTS")

    out = Path(__file__).parent / "vocence-clean.docx"
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
